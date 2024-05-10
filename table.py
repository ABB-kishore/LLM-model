# Import necessary libraries
import json
import pandas as pd
from docx import Document
import docx
from rag import results
import re

# Function to fill the table
def fill_table(tab_prompt):
    # Get the user input for table data
    tab_ans = results(tab_prompt)
    
    # Extract the JSON content from the user input
    pattern = r'```json(.*?)```'
    match = re.search(pattern, tab_ans, re.DOTALL)
    if match:
        json_content = match.group(1)
    
    # Convert the JSON content to a dictionary
    data_dict = json.loads(json_content)
    
    # Flatten the data dictionary
    flattened_data = []
    for area, area_data in data_dict['areas'].items():
        for system_type, systems in area_data['systems'].items():
            system_data = {'Area': area, 'System Type': system_type}
            for system, values in systems.items():
                if isinstance(values, dict):
                    for key, value in values.items():
                        if key in system_data:
                            system_data[key] += value
                        else:
                            system_data[key] = value
                else:
                    system_data[system] = values
            flattened_data.append(system_data)
    
    # Create a DataFrame from the flattened data
    df = pd.DataFrame(flattened_data)
    
    # Add the table to the document
    add_table(df)

# Function to add the table to the document
def add_table(df):
    # Get the table data
    bom_table = df
    
    # Load the document
    doc = Document("iminput.docx")
    
    # Find the placeholder in the document and replace it with the table
    for paragraph in doc.paragraphs:
        if "[table]" in paragraph.text:
            # Create a new table with the same number of columns as the DataFrame
            table = doc.add_table(rows=1, cols=len(bom_table.columns))
            table.style = 'Table Grid'
            
            # Add table headers
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(bom_table.columns):
                hdr_cells[i].text = header
            
            # Add table rows
            for _, row in bom_table.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
            
            # Replace the placeholder with the table in the document
            for i, element in enumerate(doc.element.body):
                if isinstance(element, docx.oxml.table.CT_Tbl):
                    first_cell_text = element.tr_lst[0].tc_lst[0].p_lst[0].text
                    if first_cell_text == "[table]":
                        doc.element.body[i] = table._tbl
                        break
    
    # Save the modified document
    doc.save("tabinput.docx")
