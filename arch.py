# Import the required modules
import shutil
import fitz
from docx import Document
import os
from docx.shared import Cm
from app import pdf_path
from rag import results
import streamlit as st
import arch
import re
from pathlib import Path

# Define the image function
def image(page):
    # Create a Path object for the save directory
    save_dir = Path("pdf_images")
    
    # Remove the save directory if it exists
    if save_dir.exists():
        shutil.rmtree(save_dir)
    
    # Create the save directory
    save_dir.mkdir(parents=True, exist_ok=True)
    
    # Open the PDF file
    with fitz.open(pdf_path) as pdf:
        p = pdf[page - 1]
        image = p.get_pixmap()
        image.save(str(save_dir / f"page_{page}.png"))
    
    # Set the image directory
    image_dir = 'pdf_images'
    
    # Open the input document
    doc = Document('input.docx')
    
    # Iterate over the tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '[image]' in cell.text:
                    cell.text = ''
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    for image_file in os.listdir(image_dir):
                        if image_file.lower().endswith(('.png')):
                            image_path = os.path.join(image_dir, image_file)
                            run.add_picture(image_path, width=Cm(18), height=Cm(10))
                            break  
 
    # Save the modified document
    doc.save('iminput.docx')


# Define the architec function
def architec():
    # Get the page number from the user
    page = results("Give me the page no. of 'DOW for Control System' from the index in the document. give me the answer in number format only.")
    page = int(re.search(r'\d+', page).group())
    
    # Display the page number
    st.write("Architecture is in page "+ str(page)+ " .")
    
    # Call the image function
    arch.image(page)

